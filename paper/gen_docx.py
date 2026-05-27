"""Build the SBLP 2026 short-paper draft as a .docx for review.

Produces an academic-style Word document mirroring the revised main.tex
content. Intended for reviewers who prefer Word over LaTeX (e.g., the
first review pass by a co-author or external reader).

Run from the repo root:
    source .venv/bin/activate
    python paper/gen_docx.py

Output: paper/Epi-SBLP2026-draft.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm


OUTPUT = Path(__file__).parent / "Epi-SBLP2026-draft.docx"


# --------------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------------


def set_cell_border(cell, **edges):
    """Set table cell borders. Each kwarg is one of 'top','bottom','left','right'
    with value like {'sz': '4', 'val': 'single', 'color': '000000'}."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge, attrs in edges.items():
        elem = OxmlElement(f'w:{edge}')
        for k, v in attrs.items():
            elem.set(qn(f'w:{k}'), str(v))
        tc_borders.append(elem)
    tc_pr.append(tc_borders)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x10, 0x20, 0x30)
    return h


def para(doc, text=None, italic=False, bold=False, justified=True, size_pt=11):
    p = doc.add_paragraph()
    if justified:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text is None:
        return p
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.italic = italic
    run.bold = bold
    return p


def add_runs(p, segments):
    """segments: list of (text, kwargs) where kwargs may include bold, italic, mono, size."""
    for seg in segments:
        if isinstance(seg, str):
            run = p.add_run(seg)
            run.font.size = Pt(11)
            continue
        text, opts = seg
        run = p.add_run(text)
        run.font.size = Pt(opts.get('size', 11))
        if opts.get('bold'):
            run.bold = True
        if opts.get('italic'):
            run.italic = True
        if opts.get('mono'):
            run.font.name = 'Consolas'
            run.font.size = Pt(opts.get('size', 10))
    return p


def code_block(doc, lines, size_pt=9):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run('\n')
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(size_pt)
        run.font.color.rgb = RGBColor(0x10, 0x20, 0x30)
    return p


# --------------------------------------------------------------------------
# Document construction
# --------------------------------------------------------------------------


doc = Document()

# Default margins (slightly narrower for a denser academic feel)
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Set base font to Times-like
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)


# --- HEADER: draft notice ---
header_p = doc.add_paragraph()
header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = header_p.add_run('DRAFT — for review by Prof. Eliseo Reátegui')
hr.italic = True
hr.font.size = Pt(9)
hr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

draft_meta = doc.add_paragraph()
draft_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
mr = draft_meta.add_run('Target venue: SBLP 2026 — Short Paper (4 pages excluding refs)')
mr.italic = True
mr.font.size = Pt(9)
mr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# --- TITLE ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run(
    'Epi: An Epistemic Type System for Containing\n'
    'LLM Hallucination in Generated Code'
)
tr.bold = True
tr.font.size = Pt(18)

# --- AUTHORS ---
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
ar = authors.add_run('Randerson Rebouças    ·    Dante Barone')
ar.font.size = Pt(12)

aff = doc.add_paragraph()
aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
afr = aff.add_run(
    'Postgraduate Program in Informatics in Education (PPGIE)\n'
    'Universidade Federal do Rio Grande do Sul (UFRGS), Porto Alegre, Brazil'
)
afr.italic = True
afr.font.size = Pt(10)

emails = doc.add_paragraph()
emails.alignment = WD_ALIGN_PARAGRAPH.CENTER
er = emails.add_run('randerson.melville@gmail.com    ·    barone@inf.ufrgs.br')
er.font.size = Pt(10)
er.font.color.rgb = RGBColor(0x33, 0x33, 0x66)


# --- ABSTRACT ---
heading(doc, 'Abstract', level=2)
para(doc,
    'We introduce Epi, a domain-specific language whose type system makes the boundary '
    'between deterministic computation and LLM-inferred values both explicit and enforceable '
    'at compile time. From a single .epi source — entities, authorization rules, AI inference '
    'steps, orchestration pipelines — the Epi transpiler emits a complete Next.js project: '
    'database schema, API routes, authorization middleware, runtime validators, LLM call '
    'wrappers, and human-review checkpoints. Every AI-inferred field carries a type-level '
    'contract — an enum, a numeric range, a confidence threshold, an optional Bayesian prior '
    '— and the transpiler discharges that contract in the generated code. The developer does '
    'not write the validation: the type does. We argue that LLM hallucination cannot be '
    'prevented, but it can be contained — by construction, at a type-system boundary in '
    'generated code. We exercise Epi end-to-end on the same source compiled against '
    'Anthropic Claude (cloud) and Qwen 2.5 3B (local, via Ollama). On an educational '
    'assessment task, the boundary routes a low-confidence misclassification to teacher '
    'review — a behavior that follows structurally from one type annotation, not from '
    'handwritten code. Epi is open-source under Apache 2.0.'
)

kw = doc.add_paragraph()
kw_run = kw.add_run('Keywords: ')
kw_run.bold = True
kw_run.font.size = Pt(10)
kw_run2 = kw.add_run(
    'Type Systems · Domain-Specific Languages · Large Language Models · '
    'Code Generation · Information-Flow Control'
)
kw_run2.font.size = Pt(10)


# --- §1 INTRODUCTION ---
heading(doc, '1. Introduction')

para(doc,
    'Large Language Models (LLMs) are increasingly embedded in production software, '
    'performing classification, extraction, summarization, and generation alongside '
    'conventional business logic. Yet the languages we use to build that software '
    'were designed for an era of deterministic computation: a function, given the '
    'same input, returns the same output. This assumption silently fails the moment '
    'an LLM enters the call graph.'
)
para(doc,
    'We identify three concrete problems that arise when building LLM-augmented '
    'applications in conventional languages.'
)

# P1
p = para(doc)
add_runs(p, [
    ('P1 — No type-level separation of epistemic domains. ', {'bold': True}),
    'A database column declared ',
    ('status: String', {'mono': True}),
    ' and an LLM-inferred column declared ',
    ('risk: String', {'mono': True}),
    ' are treated identically by the host type system. If the LLM hallucinates ',
    ('"Unknown"', {'mono': True}),
    ' in place of an expected label, the value reaches the database with no '
    'type-level alarm.',
])

p = para(doc)
add_runs(p, [
    ('P2 — Forced stack fragmentation. ', {'bold': True}),
    'A single AI-augmented feature forces the developer to hold four mental '
    'models at once: SQL or Prisma for the database, TypeScript or Python for '
    'the API, React for the UI, and a vendor SDK for the LLM. One conceptual '
    'decision is distributed across four grammars and four toolchains.',
])

p = para(doc)
add_runs(p, [
    ('P3 — No grammar-enforced hallucination contract. ', {'bold': True}),
    'Output validation, retry strategies, confidence thresholding, fallback '
    'dispatch, and audit trails are all ',
    ('optional', {'italic': True}),
    ' features that the developer must remember to add. In practice, all five '
    'are routinely omitted.',
])

para(doc,
    'This paper introduces Epi (Epistemic Programming Interface), a '
    'domain-specific language and transpiler that addresses these problems '
    'together. Our contributions are:'
)

contributions = [
    ('The Epistemic Type System', ', a type discipline that separates rigid (deterministic) and epistemic (AI-inferred) values at the type level and requires runtime validation at the boundary by construction.'),
    ('The Epi language', ', with five orthogonal primitives (Entity, Guard, Pulse, Pipeline, Lens) and a Lark/EBNF grammar; a fragment of the grammar appears in Section 3.'),
    ('The Epi transpiler', ', a three-layer architecture in which the LLM is formally excluded from the deterministic layers and operates only within constraints emitted by them.'),
    ('An end-to-end worked example', ' showing that the same .epi program runs against a frontier cloud LLM and a small local open-weight model, and that the type system catches a low-confidence misclassification as the type contract requires — presented as a concrete demonstration, not a controlled study.'),
]
for i, (head, body) in enumerate(contributions, start=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.add_run(f'{i}. ').bold = True
    r = p.add_run(head)
    r.bold = True
    p.add_run(body)

p = para(doc)
add_runs(p, [
    'Our thesis is narrow and we state it explicitly: ',
    ('LLM hallucination cannot be prevented, but it can be contained — '
     'by construction, at a type-system boundary in generated code.',
     {'bold': True}),
])


# --- §2 BACKGROUND AND RELATED WORK ---
heading(doc, '2. Background and Related Work')

para(doc, 'Epi sits at the intersection of three lines of work.')

p = para(doc)
add_runs(p, [
    ('Information-flow control. ', {'bold': True}),
    'The Rigid/Epistemic distinction is structurally analogous to the separation '
    'of security levels in information-flow analyses. Russo and Sabelfeld [11] '
    'contrast dynamic and static flow-sensitive analyses, both of which encode '
    'non-interference between value domains. Epi adapts this discipline to ',
    ('epistemic provenance', {'italic': True}),
    ' — whether a value originated from deterministic code or from probabilistic '
    'inference — rather than to confidentiality levels. The transpiler enforces '
    'non-interference by emitting validation at every domain crossing.',
])

p = para(doc)
add_runs(p, [
    ('Probabilistic programming. ', {'bold': True}),
    'Probabilistic programming languages such as ProbZelus [2] and SlicStan [4] '
    'embed distributions as first-class constructs and distinguish deterministic '
    'from probabilistic computation. Epi draws on this lineage but operates one '
    'layer higher: it orchestrates pre-trained model endpoints rather than '
    'implementing inference algorithms.',
])

p = para(doc)
add_runs(p, [
    ('Typed LLM signatures and full-stack DSLs. ', {'bold': True}),
    'BAML [3] provides typed signatures for LLM functions, and DSPy [5] compiles '
    'declarative signatures into optimized prompts. Both operate at the granularity '
    'of a single LLM call. Epi lifts typed signatures to ',
    ('whole-application generation', {'italic': True}),
    ': database, middleware, validators, and UI scaffolding emitted from one source. '
    'Wasp [13] is a full-stack DSL targeting React and Node, but it has no '
    'type-level distinction for AI-inferred values. What separates Epi from all '
    'three is the epistemic type discipline itself: audit-by-construction is a '
    'property of the type, not a feature the developer can forget to add.',
])


# --- §3 EPISTEMIC TYPE SYSTEM ---
heading(doc, '3. The Epistemic Type System')

para(doc, 'In Epi, every value belongs to one of two type domains.')

p = para(doc)
add_runs(p, [
    ('Rigid types', {'bold': True}),
    ' are deterministic and have no AI involvement:',
])

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('UUID  ·  Text  ·  Int  ·  Float  ·  Decimal  ·  Bool  ·  DateTime  ·  JSON')
r.font.name = 'Consolas'
r.font.size = Pt(10)

p = para(doc)
add_runs(p, [
    ('Epistemic types', {'bold': True}),
    ' are AI-inferred and carry mandatory validation parameters: ',
    ('AI.Enum, AI.Text, AI.Classification, AI.Score', {'mono': True}),
    ', and ',
    ('AI.Embedding', {'mono': True}),
    ', each parameterized by its specific constraints (allowed values, max tokens, '
    'dimensions, etc.) plus optional cross-cutting attributes such as ',
    ('strict', {'mono': True}),
    ', ',
    ('prior', {'mono': True}),
    ', and ',
    ('confidence_threshold', {'mono': True}),
    '.',
])

para(doc, 'The two domains are also syntactically distinct in the grammar. The '
          'relevant EBNF productions are:')

code_block(doc, [
    'entity_decl    ::= "Entity" IDENT "{" field_list "}"',
    'field_decl     ::= IDENT ":" type',
    'type           ::= rigid_type | epistemic_type',
    'rigid_type     ::= "UUID" | "Text" | "Int" | "Float"',
    '                 | "Decimal" | "Bool" | "DateTime"',
    '                 | "JSON"',
    'epistemic_type ::= "AI.Enum"   "(" enum_args   ")"',
    '                 | "AI.Text"   "(" text_args   ")"',
    '                 | "AI.Score"  "(" score_args  ")"',
    '                 | "AI.Embedding" "(" emb_args ")"',
    'enum_args      ::= ident_list ("," enum_attr)*',
    'enum_attr      ::= "strict" ":" bool',
    '                 | "prior" ":" distribution',
    '                 | "confidence_threshold" ":" number',
    'pulse_decl     ::= "Pulse" IDENT "{" pulse_body "}"',
])

para(doc, 'A field declared with an epistemic type imposes obligations on the '
          'generated code:')

code_block(doc, [
    'risco: AI.Enum(Alto, Medio, Baixo,',
    '               strict: true,',
    '               confidence_threshold: 0.85)',
])

para(doc, 'This single declaration requires the transpiler to emit, in the output '
          'project, four artifacts:')

artifacts = [
    ('A database column', ' of compatible storage type.'),
    ('A runtime validator', ' (a Zod schema in TypeScript output) that rejects values outside the declared enum.'),
    ('An LLM inference call', ' wrapped to return JSON containing the value and a _confidence field in [0,1].'),
    ('A confidence-aware dispatch', ': if reported confidence is below the threshold, route to a Checkpoint endpoint for human review; otherwise validate and persist.'),
]
for i, (head, body) in enumerate(artifacts, start=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.add_run(f'{i}. ').bold = True
    r = p.add_run(head)
    r.bold = True
    p.add_run(body)

p = para(doc)
add_runs(p, [
    'None of these artifacts is optional: a type-correct Epi program cannot omit '
    'any of them, because the type itself forces their emission. This is the '
    'central design choice: the ',
    ('type carries an obligation that the transpiler discharges in the generated '
     'code', {'italic': True}),
    '. In Pierce\'s terms [9], the transpiler refuses to emit programs that '
    'persist unvalidated AI output. The invariant we maintain informally is the '
    'following: for every well-typed Epi program, every path from an ',
    ('AI.*', {'mono': True}),
    '-typed expression to a persistence sink in the generated code passes through '
    'a validator on that type. Section 6 discusses why we describe this invariant '
    'as informal in the current stage.',
])

para(doc,
    'Crucially, typed-signature DSLs such as BAML and DSPy generate only artifact '
    '(3) above; Epi additionally generates (1), (2), and (4) as structural '
    'consequences of the same declaration.'
)

p = para(doc)
add_runs(p, [
    ('The five primitives. ', {'bold': True}),
    'Programs are composed of Entity declarations (typed schemas), Guard '
    'declarations (authorization predicates over the request context), Pulse '
    'declarations (AI execution units with prompt files, temperature, fallback '
    'strategies), Pipeline declarations (composed Pulse flows with error '
    'policies), and Lens declarations (UI surface descriptors). Of these, only '
    'Pulse and the epistemic fields within Entity invoke the AI; the others are '
    'deterministic.',
])

p = para(doc)
add_runs(p, [
    ('Trace and Checkpoint. ', {'bold': True}),
    'A Pulse may decompose into a sequence of ',
    ('Trace', {'italic': True}),
    ' steps. Each Trace can ',
    ('Expose', {'mono': True}),
    ' intermediate reasoning fields and pause at a ',
    ('Checkpoint', {'mono': True}),
    ' for human review before the next step proceeds. The transpiler emits an '
    'inspect/resume HTTP API and a persistent audit trail of each approval or '
    'correction. The mechanism targets high-stakes flows — educational '
    'assessment, clinical pre-screening, legal triage — where the reasoning '
    'must be auditable, not only the outcome.',
])

p = para(doc)
add_runs(p, [
    ('Bayesian update. ', {'bold': True}),
    'An ',
    ('AI.Enum', {'mono': True}),
    ' field may declare a prior distribution over its labels. The transpiler '
    'emits a Bayesian-update function that combines the prior with the model\'s '
    'reported likelihoods via Bayes\' rule. The posterior couples model output '
    'with domain base rates rather than overwriting them.',
])

para(doc,
    'We now describe how a Lark/EBNF grammar and these typing constraints become '
    'a runnable Next.js project.'
)


# --- §4 TRANSPILER ARCHITECTURE ---
heading(doc, '4. Transpiler Architecture')

para(doc,
    'The Epi transpiler is organized in three layers. Layer 1 parses the .epi '
    'source via a Lark grammar into a typed AST. Layer 2, the Rigid Generator, '
    'emits all deterministic artifacts: Prisma schema, Next.js routes, '
    'authorization middleware, and Zod validators. Layer 3, the Epistemic '
    'Generator, emits LLM call wrappers, Trace and Checkpoint infrastructure, '
    'and Bayesian update functions. One .epi source therefore replaces the four '
    'grammars enumerated in P2.'
)

para(doc,
    'The LLM is formally excluded from Layers 1 and 2; the deterministic layer '
    'encodes the constraints, and the epistemic layer can only operate within '
    'them. The structural similarity to the standard library / user code split '
    'in mainstream type-system implementations is intentional: the constraints '
    'are not a runtime check bolted on afterward, but a compile-time invariant '
    'the runtime cannot violate without failing closed.'
)

p = para(doc)
add_runs(p, [
    ('Provider-agnostic LLM client. ', {'bold': True}),
    'Layer 3 emits a single file, ',
    ('lib/llm-client.ts', {'mono': True}),
    ', that abstracts the LLM provider behind a uniform ',
    ('llmCall(params)', {'mono': True}),
    ' interface. Runtime environment variables (',
    ('EPI_AI_PROVIDER, EPI_AI_MODEL, EPI_AI_BASE_URL', {'mono': True}),
    ') select between Anthropic Claude [1] and any OpenAI-compatible endpoint; '
    'we target Ollama [7] for local open-weight models. The epistemic contract '
    '— JSON output containing a ',
    ('_confidence', {'mono': True}),
    ' field, parseable by the Zod schema — is enforced uniformly across '
    'providers. Provider choice is a deployment concern, not a code change.',
])


# --- §5 DEMONSTRATION ---
heading(doc, '5. Demonstration')

para(doc,
    'We exercise the system end-to-end on a single educational scenario, '
    'presented as a worked example rather than a controlled study; a controlled '
    'comparison against hand-written baselines on correctness coverage, audit '
    'completeness, and developer time is planned for the next phase.'
)

p = para(doc)
add_runs(p, [
    'The .epi source declares one Entity ',
    ('Submissao', {'mono': True}),
    ' (student answer) with an epistemic field ',
    ('avaliacao: AI.Enum(Correto, Parcial, Incorreto, prior: Distribution('
     '0.40, 0.45, 0.15), confidence_threshold: 0.85)', {'mono': True, 'size': 9}),
    ', one Guard for teacher authentication, one Pulse ',
    ('AvaliarResposta', {'mono': True}),
    ' invoking ',
    ('AI.classify', {'mono': True}),
    ' with ',
    ('on_low_confidence: Checkpoint.ReviewRequired(role: "Professor")', {'mono': True, 'size': 9}),
    ', and one Pipeline. The source is 30 lines; the transpiled project is 13 '
    'files and roughly 700 lines of TypeScript and Prisma schema.',
])

p = para(doc)
add_runs(p, [
    'We compile the same .epi source twice and run both deployments unchanged: '
    '(a) against ',
    ('Anthropic Claude Sonnet 4', {'bold': True}),
    ' in the cloud; (b) against ',
    ('Qwen 2.5 3B Instruct', {'bold': True}),
    ' (Q4_K_M, ~1.9 GB) served locally by Ollama on consumer hardware (Apple '
    'Silicon, 16 GB RAM). The only difference is the value of ',
    ('EPI_AI_PROVIDER', {'mono': True}),
    ' in the deployment\'s .env file. Each deployment receives the same four '
    'inputs against the prompt "What is photosynthesis?". Results appear in '
    'Table 1.',
])

# --- TABLE 1 (evaluation) ---
para(doc, 'Table 1. Four inputs exercised end-to-end. The fourth row injects a '
          'deliberately malformed JSON output to exercise the schema-validator '
          'path. Latencies measured wall-clock on Apple Silicon (16 GB RAM).',
     italic=True, justified=False, size_pt=9)

table = doc.add_table(rows=5, cols=5)
table.alignment = WD_ALIGN_PARAGRAPH.CENTER
table.autofit = True

# Headers
headers = ['Student answer', 'Qwen 2.5 3B output', 'Claude Sonnet 4 output',
           'Boundary action', 'Lat.']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(9)
    set_cell_border(cell,
        top={'sz': '8', 'val': 'single', 'color': '000000'},
        bottom={'sz': '6', 'val': 'single', 'color': '000000'},
    )

rows_data = [
    [
        'Textbook-correct full definition',
        '{"avaliacao":"Parcial", "_confidence":0.8}',
        '{"avaliacao":"Correto", "_confidence":0.93}',
        'Qwen: 0.8 < 0.85 → Checkpoint to Professor. Claude: validated → "Correto" persisted.',
        '2.6s / 1.1s',
    ],
    [
        '"Plants eat sunlight."',
        '{"avaliacao":"Incorreto", "_confidence":0.9}',
        '{"avaliacao":"Incorreto", "_confidence":0.95}',
        'Both validated → "Incorreto" persisted.',
        '0.8s / 0.9s',
    ],
    [
        '"Photosynthesis is my dog\'s name."',
        '{"avaliacao":"Incorreto", "_confidence":0.9}',
        '{"avaliacao":"Incorreto", "_confidence":0.97}',
        'Both validated → "Incorreto" persisted.',
        '0.8s / 0.8s',
    ],
    [
        'Induced malformed model output (control)',
        '{"unexpected":"object"}',
        '—',
        'Zod rejects → Fallback.ManualReview queue.',
        '—',
    ],
]
for i, row in enumerate(rows_data, start=1):
    for j, cell_text in enumerate(row):
        cell = table.rows[i].cells[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        r = p.add_run(cell_text)
        # Monospace for JSON cells
        if j in (1, 2) and (cell_text.startswith('{') or cell_text == '—'):
            r.font.name = 'Consolas'
            r.font.size = Pt(8)
        else:
            r.font.size = Pt(9)

# Bottom border on last row
for cell in table.rows[-1].cells:
    set_cell_border(cell, bottom={'sz': '8', 'val': 'single', 'color': '000000'})


# --- §5 Observation paragraphs ---
p = para(doc)
add_runs(p, [
    ('Observation. ', {'bold': True}),
    'The first row is the contribution of this evaluation. The small model '
    'misclassified a textbook-correct answer as "Parcial" — a false negative '
    'that, if persisted, would unfairly penalize the student. The model\'s own '
    'reported confidence — 0.8 — fell below the type-declared threshold of '
    '0.85. The transpiler-emitted dispatch routed the case to a human reviewer ',
    ('before', {'italic': True}),
    ' it reached the database. ',
    ('The developer wrote no code for this behavior.', {'bold': True}),
    ' It follows from the single annotation ',
    ('confidence_threshold: 0.85', {'mono': True}),
    ' on the type, which the transpiler discharged in Layer 3. By contrast, '
    'BAML and DSPy would type-validate the model output; only Epi additionally '
    'emits the Checkpoint dispatch and the audit trail without developer code.',
])

p = para(doc)
add_runs(p, [
    ('Provider parity. ', {'bold': True}),
    'Submitting the same three real inputs to Anthropic Claude Sonnet 4 '
    'classified the textbook-correct answer as "Correto" with confidence above '
    '0.9, validating directly. The pipeline behaves identically at the boundary '
    'across both providers; what differs is absolute accuracy and latency. This '
    'confirms that the epistemic contract is enforced independently of the '
    'inference engine.',
])

p = para(doc)
add_runs(p, [
    ('Hallucination containment. ', {'bold': True}),
    'The fourth row of Table 1 exercises the schema-validator path with a '
    'deliberately malformed JSON. The Zod schema rejects the value and the '
    'transpiler-emitted ',
    ('Fallback.ManualReview', {'mono': True}),
    ' routes the request to a manual queue. Once more, no handwritten error '
    'path was required.',
])

para(doc,
    'The demonstration above shows the mechanism on one scenario; the section '
    'that follows marks the boundaries of what it does not yet establish.'
)


# --- §6 DISCUSSION AND LIMITATIONS ---
heading(doc, '6. Discussion and Limitations')

p = para(doc)
add_runs(p, [
    ('Type-system formalization. ', {'bold': True}),
    'The current type-system specification is informal: we present a working '
    'transpiler and an empirical demonstration, not a mechanized soundness '
    'proof. The invariant we describe in Section 3 — ',
    ('every path from an AI.*-typed expression to a persistence sink passes '
     'through a validator on that type', {'italic': True}),
    ' — is a non-interference property between the rigid and epistemic value '
    'spaces. A mechanized treatment in Coq or Agda, including a typing '
    'judgment and a soundness theorem, is future work.',
])

p = para(doc)
add_runs(p, [
    ('Coverage of targets. ', {'bold': True}),
    'The transpiler currently emits a complete Next.js project. A FastAPI '
    'target is partially implemented and gated in the CLI until it is '
    'complete. This is an engineering limitation, not an essential property '
    'of the design.',
])

p = para(doc)
add_runs(p, [
    ('Lens.Mood. ', {'bold': True}),
    'The Lens primitive includes a ',
    ('Mood', {'mono': True}),
    ' declaration that currently maps keywords to UI utility classes via a '
    'deterministic lookup table. It is the weakest part of the current '
    'design and is marked experimental in the documentation.',
])


# --- §7 CONCLUSION ---
heading(doc, '7. Conclusion and Future Work')

p = para(doc)
add_runs(p, [
    'LLM hallucination is a property of the underlying inference, not of the '
    'surrounding code; it cannot be prevented at the type-system level. It '
    'can, however, be ',
    ('contained', {'italic': True}),
    '. A type discipline that distinguishes deterministic from epistemic '
    'provenance can require — by construction, in generated code — that '
    'every AI-inferred value pass a validator, report its confidence, and '
    'surface for human review when uncertain. We have implemented this '
    'discipline as a language and a transpiler, exercised it end-to-end '
    'against a frontier cloud model and a small local open-weight model, '
    'and shown that the boundary catches a misclassification that would '
    'otherwise reach a database of record. Three directions follow: a '
    'mechanized formalization of the type system, a controlled comparison '
    'against hand-written baselines on correctness and audit coverage, and '
    'a classroom study of Trace+Checkpoint adoption in educational '
    'assessment.',
])


# --- ARTIFACT AVAILABILITY ---
heading(doc, 'Artifact Availability', level=2)

p = para(doc)
add_runs(p, [
    'Epi is open-source software under the Apache 2.0 license. The full '
    'source — language, transpiler, grammar, parser, tests, the ',
    ('avaliacao-simples.epi', {'mono': True}),
    ' program exercised in Section 5, and the generated TypeScript artifacts '
    '— is publicly available. The PyPI distribution can be installed with ',
    ('pip install epi-lang', {'mono': True}),
    '. A landing page summarizes the project, its scope, and known '
    'limitations. (Repository URL and landing page URL are omitted in this '
    'anonymized review draft and will be added in the camera-ready version.)',
])


# --- REFERENCES ---
heading(doc, 'References', level=2)

refs = [
    'Anthropic. Claude API Documentation. https://docs.anthropic.com, 2025.',
    'Baudart, G., Mandel, L., Atkinson, E., Sherman, B., Pouzet, M., Carbin, M. Reactive Probabilistic Programming. PLDI 2020, pp. 898–912. DOI 10.1145/3385412.3386009.',
    'BoundaryML. BAML: A Domain-Specific Language for AI Engineering. https://github.com/BoundaryML/baml, 2024.',
    'Gorinova, M. I., Gordon, A. D., Sutton, C. Probabilistic Programming with Densities in SlicStan: Efficient, Flexible, and Deterministic. POPL 2019, vol. 3. DOI 10.1145/3290348.',
    'Khattab, O., Singhvi, A., Maheshwari, P., et al. DSPy: Compiling Declarative Language Model Calls Into Self-Improving Pipelines. arXiv:2310.03714, 2023.',
    'Ollama Inc. Ollama: Run Open-Weight LLMs Locally. https://ollama.com, 2024.',
    'Pierce, B. C. Types and Programming Languages. MIT Press, 2002.',
    'Russo, A., Sabelfeld, A. Dynamic vs. Static Flow-Sensitive Security Analysis. CSF 2010, pp. 186–199. DOI 10.1109/CSF.2010.20.',
    'Wasp. Wasp: A DSL for Full-Stack Web Apps. https://wasp-lang.dev, 2024.',
]
for i, ref in enumerate(refs, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.7)
    run_num = p.add_run(f'[{i}]  ')
    run_num.bold = True
    run_num.font.size = Pt(10)
    r = p.add_run(ref)
    r.font.size = Pt(10)


# --- SAVE ---
doc.save(str(OUTPUT))
print(f'Wrote {OUTPUT}')
print(f'Size: {OUTPUT.stat().st_size / 1024:.1f} KB')
